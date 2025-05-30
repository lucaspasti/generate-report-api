import os
import io
from uuid import UUID
from collections import OrderedDict
import datetime
from datetime import datetime
import numpy as np
import pandas as pd
import requests
from pytz import timezone
from supabase import Client
from fastapi import HTTPException
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
from app.schemas.qag import QAGRequest, QAGResponse
from app.utils.date_utils import mes_por_extenso
from app.utils.file_utils import create_signed_url
from app.utils.graficos import grafico_qualidade_agua
from app.services.indicadores.indicadores_qag import indicadores_qag
from app.services.vmps.vmp_qag import vmp_qag

# Caminho para o template DOCX
TEMPLATE_PATH = os.path.join(
    os.getcwd(), "app/services/relatorios/qualidade_agua_superficial_u.docx"
)
BUCKET = "relatorios-qag"

# Parâmetros pré-definidos (podem ser importados de outro módulo se preferir)
parametros_fisico_quimicos = [
    "Materiais flutuantes",
    "Óleos e graxas",
    "Substâncias que comuniquem gosto ou odor",
    "Corantes provenientes de fontes antrópicas",
    "Resíduos sólidos objetáveis",
    "Turbidez (UNT)",
    "Cor verdadeira (mg Pt/L)",
    "Sólidos dissolvidos totais (mg/L)",
    "pH (N/A)",
]
# --- Parâmetros pré-definidos ---
parametros_metais_pesados = [
    "Antimônio (mg/L Sb)",
    "Arsênio total (mg/L As)",
    "Bário total (mg/L Ba)",
    "Berílio total (mg/L Be)",
    "Boro total (mg/L B)",
    "Cádmio total (mg/L Cd)",
    "Chumbo total (mg/L Pb)",
    "Cobalto total (mg/L Co)",
    "Cobre dissolvido (mg/L Cu)",
    "Cromo total (mg/L Cr)",
    "Ferro dissolvido (mg/L Fe)",
    "Manganês total (mg/L Mn)",
    "Mercúrio total (mg/L Hg)",
    "Níquel total (mg/L Ni)",
    "Prata total (mg/L Ag)",
    "Selênio total (mg/L Se)",
    "Tálio total (mg/L Tl)",
    "Urânio total (mg/L U)",
    "Vanádio total (mg/L V)",
    "Zinco total (mg/L Zn)",
]

parametros_fisico_quimicos = [
    "Materiais flutuantes",
    "Óleos e graxas",
    "Substâncias que comuniquem gosto ou odor",
    "Corantes provenientes de fontes antrópicas",
    "Resíduos sólidos objetáveis",
    "Turbidez (UNT)",
    "Cor verdadeira (mg Pt/L)",
    "Sólidos dissolvidos totais (mg/L)",
    "pH (N/A)",
]

# Oxigênio
parametros_oxigenio = ["DBO 5 dias a 20°C (mg/L O2)", "OD (mg/L O2)"]

# Microbiológicos
parametros_microbiologicos = ["Coliformes termotolerantes (NMP/100mL)"]

# Nutrientes
nutrientes = [
    "Fósforo total (ambiente lêntico) (mg/L P)",
    "Fósforo total (ambiente intermediário) (mg/L P)",
    "Fósforo total (ambiente lótico) (mg/L P)",
    "Polifosfatos (mg/L P)",
    "Nitrato (mg/L N)",
    "Nitrito (mg/L N)",
    "Nitrogênio amoniacal total (pH ≤ 7,5) (mg/L N)",
    "Nitrogênio amoniacal total (7,5 < pH ≤ 8,0) (mg/L N)",
    "Nitrogênio amoniacal total (8,0 < pH ≤ 8,5) (mg/L N)",
    "Nitrogênio amoniacal total (pH > 8,5) (mg/L N)",
]

# Outros elementos dissolvidos
elementos_dissolvidos = [
    "Alumínio dissolvido (mg/L Al)",
    "Antimônio (mg/L Sb)",
    "Arsênio total (mg/L As)",
    "Bário total (mg/L Ba)",
    "Berílio total (mg/L Be)",
    "Boro total (mg/L B)",
    "Chumbo total (mg/L Pb)",
    "Cobalto total (mg/L Co)",
    "Cobre dissolvido (mg/L Cu)",
    "Cromo total (mg/L Cr)",
    "Ferro dissolvido (mg/L Fe)",
    "Fluoreto total (mg/L F)",
    "Lítio total (mg/L Li)",
    "Manganês total (mg/L Mn)",
    "Mercúrio total (mg/L Hg)",
    "Níquel total (mg/L Ni)",
    "Prata total (mg/L Ag)",
    "Selênio total (mg/L Se)",
    "Tálio total (mg/L Tl)",
    "Urânio total (mg/L U)",
    "Vanádio total (mg/L V)",
    "Zinco total (mg/L Zn)",
]

# PAHs (Hidrocarbonetos aromáticos policíclicos)
pahs = [
    "Benzo (a) antraceno (μg/L)",
    "Benzo (a) pireno (μg/L)",
    "Benzo(b) fluoranteno (μg/L)",
    "Benzo(k) fluoranteno (μg/L)",
    "Criseno (μg/L)",
    "Dibenzo (a,h) antraceno (μg/L)",
]

# Pesticidas e PCBs
pesticidas_pcbs = [
    "Acrilamida (μg/L)",
    "Alacloro (μg/L)",
    "Aldrin + Dieldrin (μg/L)",
    "Atrazina (μg/L)",
    "Carbaril (μg/L)",
    "Clordano (cis + trans) (μg/L)",
    "DDT (p,p'-DDT + p,p'-DDE + p,p'-DDD) (μg/L)",
    "Dodecacloro pentaciclodecano (μg/L)",
    "Endossulfan (a + b + sulfato) (μg/L)",
    "Endrin (μg/L)",
    "Lindano (g-HCH) (μg/L)",
    "Malation (μg/L)",
    "Metoxicloro (μg/L)",
    "Pentaclorofenol (μg/L)",
    "Toxafeno (μg/L)",
    "Trifluralina (μg/L)",
    "PCBs - Bifenilas Policloradas (μg/L)",
]

# Solventes halogenados e compostos voláteis
solventes = [
    "1,2-Dicloroetano (mg/L)",
    "1,1-Dicloroeteno (mg/L)",
    "2,4-D (μg/L)",
    "2,4-Diclorofenol (μg/L)",
    "2,4,5-T (μg/L)",
    "2,4,6 - Triclorofenol (mg/L)",
    "Tricloroeteno (mg/L)",
    "Tetracloreto de carbono (mg/L)",
    "Tetracloroeteno (mg/L)",
    "Diclorometano (mg/L)",
    "Estireno (mg/L)",
    "Etilbenzeno (μg/L)",
    "Tolueno (μg/L)",
    "Monoclorobenzeno (μg/L)",
    "2-Clorofenol (μg/L)",
]

# Outros orgânicos e surfactantes
outros_organicos = [
    "Fenóis totais (mg/L)",
    "Substâncias tensoativas que reagem com o azul de metileno (mg/L LAS)",
    "Substâncias tensoativas que reagem com o vermelho de metila (mg/L MBAS)",
]

# Íons comuns
outros_ions = [
    "Cloreto total (mg/L Cl)",
    "Cloro residual total (mg/L Cl)",
    "Sulfato total (mg/L SO4)",
    "Sulfeto (H2S não dissociado) (mg/L S)",
    "Carbono orgânico total (mg/L C)",
]


def gerar_relatorio_qag(supabase: Client, payload: QAGRequest) -> QAGResponse:
    """
    Gera, faz upload e registra no banco um relatório QAG.
    """
    # 1) Datas e chaves
    periodicidade = payload.periodicidade
    data_str = payload.data_campanha.isoformat()
    tz_br = timezone("America/Sao_Paulo")
    now_br = datetime.now(tz_br)
    object_key = f"{payload.ativo_id}/{data_str}/{now_br:%Y-%m-%d_%H-%M-%S}.docx"
    # use o tempdir correto do sistema
    tmp_dir = tempfile.gettempdir()
    local_path = os.path.join(tmp_dir, os.path.basename(object_key))

    # 2) Carrega template
    document = DocxTemplate(TEMPLATE_PATH)

    # 3) Consultas no Supabase
    ativo = (
        supabase.table("ativos").select("*").eq("id", str(payload.ativo_id)).execute()
    )
    if not ativo.data:
        raise HTTPException(404, detail="Ativo não encontrado")

    configuracoes = (
        supabase.table("configuracao_formulario_ativos")
        .select("*")
        .eq("ativo_id", str(payload.ativo_id))
        .eq("tipo_formulario", "form_qualidade_da_agua_superficial")
        .execute()
    )
    if not configuracoes.data:
        raise HTTPException(404, detail="Configuração do formulário não encontrada")

    form_qag = (
        supabase.table("form_qualidade_da_agua_superficial")
        .select("*")
        .eq("ativo_id", str(payload.ativo_id))
        .eq("campanha_de_coleta", data_str)
        .execute()
    )
    if not form_qag.data:
        raise HTTPException(404, detail="Campanha não encontrada")

    # 4) DataFrame de resultados
    resultados = form_qag.data[0]["resultados"]

    df_resultados = pd.DataFrame(resultados).fillna("Indisponível")

    # Parâmetros escolhidos
    parametros = parametros_fisico_quimicos

    # 5) Montagem de imagens InlineImage
    lab = form_qag.data[0]
    fotos = []
    for key in [
        "registros_fotograficos_sondas",
        "registros_fotograficos_amostradores",
        "registros_fotograficos_caixas_termicas",
    ]:
        url = lab[key][0]
        resp = requests.get(url)
        resp.raise_for_status()
        buf = io.BytesIO(resp.content)
        fotos.append(InlineImage(document, buf, width=Cm(5)))
    q_22, q_23, q_24 = fotos

    # 6) Gráficos
    selecionados = df_resultados.copy()
    selecionados = selecionados[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"] + parametros
    ]
    data_dt = datetime.strptime(data_str, "%Y-%m-%d").date()

    q_14 = df_resultados["Ponto"].nunique()

    q_20_1 = form_qag.data[0]["nome_laboratorio"]
    q_20_2 = form_qag.data[0]["razao_social_laboratorio"]
    q_20_3 = form_qag.data[0]["cnpj_laboratorio"]
    q_20_4 = form_qag.data[0]["endereco_laboratorio"]
    q_20_5 = form_qag.data[0]["responsavel_tecnico"]
    q_20_6 = form_qag.data[0]["email"]
    q_20_7 = form_qag.data[0]["contato"]

    q_25 = configuracoes.data[0]["dados_laboratoriais"][0].get("metodologia_adotada")

    q_26 = indicadores_qag

    #############################################################

    q_27 = selecionados.to_dict(orient="records")

    subdoc_27 = document.new_subdoc()
    n_cols = 2 + len(parametros)
    table = subdoc_27.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Profundidade"
    for idx, param in enumerate(parametros, start=2):
        hdr[idx].text = param

    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo a ordem de aparecimento em q_27
    grupos = OrderedDict()
    for linha in q_27:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + faz merge vertical na coluna “Ponto”
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            row[1].text = str(linha["Profundidade"])
            for j, param in enumerate(parametros, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1

        # merge das células da coluna 0 entre start_row e end_row
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto
    ################################################
    all_figs_qag = []
    for classe in selecionados["Classe"].unique():
        figs = grafico_qualidade_agua(
            selecionados[selecionados["Classe"] == classe], parametros, classe, vmp_qag
        )
        all_figs_qag.extend(figs)

    # 2) gera as InlineImage
    imagens_qag29 = []
    for fig in all_figs_qag:
        buf = io.BytesIO()
        fig.savefig(buf, format="PNG", dpi=120, bbox_inches="tight")
        buf.seek(0)
        imagem = InlineImage(document, buf, width=Cm(12), height=Cm(6))
        imagens_qag29.append(imagem)

    #####################################

    q_30 = selecionados["Profundidade"].unique()

    #####################################

    registros_31 = []

    for _, row in selecionados.iterrows():
        classe = row["Classe"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_31.append(
                {
                    "Ponto": row["Ponto"],
                    "Classe": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao = pd.DataFrame(registros_31)

    q_31 = df_comparacao["Conforme"].mean() * 100

    ####################################################

    subdoc_32 = document.new_subdoc()

    # número de linhas: 2 de cabeçalho + 1 para cada parâmetro
    n_rows = 2 + len(parametros)
    n_cols = 1 + 4  # 1 coluna Parâmetro + 4 colunas de Média

    table = subdoc_32.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr0 = table.rows[0].cells
    hdr0[0].text = "Parâmetro"
    # mescla horizontal colunas 1..4
    m = hdr0[1].merge(hdr0[2]).merge(hdr0[3]).merge(hdr0[4])
    m.text = "Média"

    # aplica bold ao hdr0
    for cell in [hdr0[0], m]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    # linha 1: labels das 4 subcolunas
    hdr1 = table.rows[1].cells
    hdr1[0].text = ""  # espaço porque Parâmetro já virou rowspan
    for idx, label in enumerate(["Total", "Superfície", "Meio", "Fundo"], start=1):
        hdr1[idx].text = label
        for p in hdr1[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True

    # --- 2) Corpo da tabela: uma linha por parâmetro ---

    # converte coluna de profundidade para facilitar filtro
    _resuldf_resultados = df_resultados.copy()
    _resuldf_resultados["Profundidade"] = _resuldf_resultados["Profundidade"].astype(
        str
    )

    for i, param in enumerate(parametros):
        row = table.rows[2 + i].cells
        row[0].text = param

        # calcula médias
        série = pd.to_numeric(_resuldf_resultados[param], errors="coerce")
        media_total = série.mean()

        media_sup = série[_resuldf_resultados["Profundidade"] == "Superfície"].mean()
        media_meio = série[_resuldf_resultados["Profundidade"] == "Meio"].mean()
        media_fundo = série[_resuldf_resultados["Profundidade"] == "Fundo"].mean()

        for col, val in zip(
            (1, 2, 3, 4), (media_total, media_sup, media_meio, media_fundo)
        ):
            row[col].text = "" if pd.isna(val) else f"{val:.2f}"

    ########################################################################

    q_33 = parametros_metais_pesados

    ########################################################################

    subdoc_34 = document.new_subdoc()

    metais_pesados = df_resultados.copy()
    metais_pesados = metais_pesados[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"]
        + parametros_metais_pesados
    ]
    q_33 = parametros_metais_pesados

    metais_pesados = metais_pesados.to_dict(orient="records")

    n_cols = 2 + len(parametros_metais_pesados)
    table = subdoc_34.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Profundidade"
    for idx, param in enumerate(parametros_metais_pesados, start=2):
        hdr[idx].text = param

    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo a ordem de aparecimento em q_27
    grupos = OrderedDict()
    for linha in metais_pesados:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + faz merge vertical na coluna “Ponto”
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            row[1].text = str(linha["Profundidade"])
            for j, param in enumerate(parametros_metais_pesados, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1

        # merge das células da coluna 0 entre start_row e end_row
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto

    ########################################################################

    registros_35 = []

    metais_pesados_35 = df_resultados.copy()
    metais_pesados_35 = metais_pesados_35[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"]
        + parametros_metais_pesados
    ]

    for _, row in metais_pesados_35.iterrows():
        classe = row["Classe"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_35.append(
                {
                    "Ponto": row["Ponto"],
                    "Classe": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao_35 = pd.DataFrame(registros_35)

    q_35 = df_comparacao_35["Conforme"].mean() * 100

    #################################################################

    metais_pesados_36 = df_resultados.copy()
    metais_pesados_36 = metais_pesados_36[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"]
        + parametros_metais_pesados
    ]

    all_figs_qag_36 = []
    for classe in metais_pesados_36["Classe"].unique():
        figs = grafico_qualidade_agua(
            metais_pesados_36[metais_pesados_36["Classe"] == classe],
            parametros_metais_pesados,
            classe,
            vmp_qag,
        )
        all_figs_qag_36.extend(figs)

    imagens_qag_36 = []
    for fig in all_figs_qag_36:
        fig.savefig(buf, format="PNG", dpi=120, bbox_inches="tight")
        buf.seek(0)
        imagem = InlineImage(document, buf, width=Cm(12), height=Cm(6))
        imagens_qag_36.append(imagem)

    ########################################################################

    subdoc_37 = document.new_subdoc()

    # número de linhas: 2 de cabeçalho + 1 para cada parâmetro
    n_rows = 2 + len(parametros_metais_pesados)
    n_cols = 1 + 4  # 1 coluna Parâmetro + 4 colunas de Média

    table = subdoc_37.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- 1) Cabeçalho em dois níveis ---

    # linha 0: "Parâmetro" (rowspan=2) + "Média" (colspan=4)
    hdr0 = table.rows[0].cells
    hdr0[0].text = "Parâmetro"
    # mescla horizontal colunas 1..4
    m = hdr0[1].merge(hdr0[2]).merge(hdr0[3]).merge(hdr0[4])
    m.text = "Média"

    # aplica bold ao hdr0
    for cell in [hdr0[0], m]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    # linha 1: labels das 4 subcolunas
    hdr1 = table.rows[1].cells
    hdr1[0].text = ""  # espaço porque Parâmetro já virou rowspan
    for idx, label in enumerate(["Total", "Superfície", "Meio", "Fundo"], start=1):
        hdr1[idx].text = label
        for p in hdr1[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True

    # --- 2) Corpo da tabela: uma linha por parâmetro ---

    # converte coluna de profundidade para facilitar filtro
    _resuldf_resultados = df_resultados.copy()
    _resuldf_resultados = _resuldf_resultados[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"]
        + parametros_metais_pesados
    ]
    _resuldf_resultados["Profundidade"] = _resuldf_resultados["Profundidade"].astype(
        str
    )

    for i, param in enumerate(parametros_metais_pesados):
        row = table.rows[2 + i].cells
        row[0].text = param

        # calcula médias
        série = pd.to_numeric(_resuldf_resultados[param], errors="coerce")
        media_total = série.mean()

        media_sup = série[_resuldf_resultados["Profundidade"] == "Superfície"].mean()
        media_meio = série[_resuldf_resultados["Profundidade"] == "Meio"].mean()
        media_fundo = série[_resuldf_resultados["Profundidade"] == "Fundo"].mean()

        for col, val in zip(
            (1, 2, 3, 4), (media_total, media_sup, media_meio, media_fundo)
        ):
            row[col].text = "" if pd.isna(val) else f"{val:.2f}"

    ########################################################################

    registros_38 = []

    organicos_38 = df_resultados.copy()
    organicos_38 = organicos_38[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"] + solventes
    ]

    for _, row in organicos_38.iterrows():
        classe = row["Classe"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_38.append(
                {
                    "Ponto": row["Ponto"],
                    "Classe": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao_38 = pd.DataFrame(registros_38)

    q_38 = df_comparacao_38["Conforme"].mean() * 100

    ###################################################################

    subdoc_39 = document.new_subdoc()

    # número de linhas: 2 de cabeçalho + 1 para cada parâmetro
    n_rows = 2 + len(solventes)
    n_cols = 1 + 4  # 1 coluna Parâmetro + 4 colunas de Média

    table = subdoc_39.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- 1) Cabeçalho em dois níveis ---

    # linha 0: "Parâmetro" (rowspan=2) + "Média" (colspan=4)
    hdr0 = table.rows[0].cells
    hdr0[0].text = "Parâmetro"
    # mescla horizontal colunas 1..4
    m = hdr0[1].merge(hdr0[2]).merge(hdr0[3]).merge(hdr0[4])
    m.text = "Média"

    # aplica bold ao hdr0
    for cell in [hdr0[0], m]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    # linha 1: labels das 4 subcolunas
    hdr1 = table.rows[1].cells
    hdr1[0].text = ""  # espaço porque Parâmetro já virou rowspan
    for idx, label in enumerate(["Total", "Superfície", "Meio", "Fundo"], start=1):
        hdr1[idx].text = label
        for p in hdr1[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True

    # --- 2) Corpo da tabela: uma linha por parâmetro ---

    # converte coluna de profundidade para facilitar filtro
    _resuldf_resultados = df_resultados.copy()
    _resuldf_resultados = _resuldf_resultados[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"] + solventes
    ]
    _resuldf_resultados["Profundidade"] = _resuldf_resultados["Profundidade"].astype(
        str
    )

    for i, param in enumerate(solventes):
        row = table.rows[2 + i].cells
        row[0].text = param

        # calcula médias
        série = pd.to_numeric(_resuldf_resultados[param], errors="coerce")
        media_total = série.mean()

        media_sup = série[_resuldf_resultados["Profundidade"] == "Superfície"].mean()
        media_meio = série[_resuldf_resultados["Profundidade"] == "Meio"].mean()
        media_fundo = série[_resuldf_resultados["Profundidade"] == "Fundo"].mean()

        for col, val in zip(
            (1, 2, 3, 4), (media_total, media_sup, media_meio, media_fundo)
        ):
            row[col].text = "" if pd.isna(val) else f"{val:.2f}"

    ########################################################################

    subdoc_40 = document.new_subdoc()

    solventes_40 = df_resultados.copy()
    solventes_40 = solventes_40[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"] + solventes
    ]

    solventes_40 = solventes_40.to_dict(orient="records")

    n_cols = 2 + len(solventes)
    table = subdoc_40.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Profundidade"
    for idx, param in enumerate(solventes, start=2):
        hdr[idx].text = param

    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo a ordem de aparecimento em q_27
    grupos = OrderedDict()
    for linha in solventes_40:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + faz merge vertical na coluna “Ponto”
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            row[1].text = str(linha["Profundidade"])
            for j, param in enumerate(solventes, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1

        # merge das células da coluna 0 entre start_row e end_row
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto

    ########################################################################

    laudo = form_qag.data[0]["laudos"][0]

    q_43 = RichText()
    q_43.add(
        laudo, underline=True, color="#1281F0", url_id=document.build_url_id(laudo)
    )

    #########################################################################

    tabela_qag47 = indicadores_qag

    tabela_qag47 = pd.DataFrame(tabela_qag47)

    tabela_qag47_aux = df_comparacao.copy()

    tabela_qag47 = tabela_qag47.merge(
        tabela_qag47_aux[["Parametro", "Valor", "Conforme"]],
        how="left",
        left_on=["Parametro"],
        right_on=["Parametro"],
    )

    tabela_qag47 = tabela_qag47.dropna()

    for _, row in tabela_qag47.iterrows():
        if row["Conforme"] == True:
            tabela_qag47.at[_, "Conforme"] = "Alcançado"
        elif row["Conforme"] == False:
            tabela_qag47.at[_, "Conforme"] = "Não Alcançado"

    tabela_qag47 = tabela_qag47.drop(
        columns=["Tipo", "Programa", "Valor", "Unidade", "Parametro"]
    )
    tabela_qag47 = tabela_qag47.rename(columns={"Conforme": "Resultado"})

    tabela_qag47 = tabela_qag47.to_dict(orient="records")

    ########################################################################

    solventes_49 = df_resultados.copy()
    solventes_49 = solventes_49[
        ["Ponto", "Classe", "Profundidade", "Tipo de análise"] + solventes
    ]

    all_figs_qag_49 = []
    for classe in solventes_49["Classe"]:
        figs = grafico_qualidade_agua(
            solventes_49[solventes_49["Classe"] == classe], solventes, classe, vmp_qag
        )
        all_figs_qag_49.extend(figs)

    imagens_qag_49 = []
    for fig in all_figs_qag_49:
        fig.savefig(buf, format="PNG", dpi=120, bbox_inches="tight")
        buf.seek(0)
        imagem = InlineImage(document, buf, width=Cm(12), height=Cm(6))
        imagens_qag_49.append(imagem)

    # 7) Contexto e renderização
    contexto = {
        "parametros_escolhidos": parametros,
        "QAG_01": ativo.data[0]["nome"],
        "QAG_02": form_qag.data[0]["campanha_de_coleta"],
        "QAG_03": data_dt.strftime("%m"),
        "QAG_04": data_dt.strftime("%Y"),
        "QAG_05": "Florianópolis",
        "QAG_06": datetime.now().day,
        "QAG_07": mes_por_extenso(data_dt.strftime("%m")),
        "QAG_08": datetime.now().year,
        "QAG_09": ativo.data[0]["nome"],
        "QAG_10": ativo.data[0]["cnpj"],
        "QAG_11": ativo.data[0]["endereco"],
        "QAG_12": ativo.data[0]["nome"],
        "QAG_13": ativo.data[0]["numero_licenca"],
        "QAG_14": q_14,
        "QAG_15": ativo.data[0]["orgao_regulador"],
        "QAG_16": ativo.data[0]["endereco"],
        #  "QAG_17": configuracoes_form.data[0]['localizacao_dos_pontos_de_monitoramento'],#pontos
        # pontos
        "QAG_18": configuracoes.data[0]["localizacao_dos_pontos_de_monitoramento"],
        "QAG_19": configuracoes.data[0]["parametro_periodicidade"],  # pontos
        "QAG_20_1": q_20_1,
        "QAG_20_2": q_20_2,
        "QAG_20_3": q_20_3,
        "QAG_20_4": q_20_4,
        "QAG_20_5": q_20_5,
        "QAG_20_6": q_20_6,
        "QAG_20_7": q_20_7,
        #        "QAG_21": configuracoes_form.data[0]['parametros_periodicidade'],#
        "QAG_22": q_22,
        "QAG_23": q_23,
        "QAG_24": q_24,
        "QAG_25": q_25,
        "QAG_26": q_26,
        "tabela_qag_27": subdoc_27,
        "QAG_28": parametros,
        "QAG_29": imagens_qag29,
        "QAG_30": q_30,
        "QAG_31": q_31,
        "QAG_32": subdoc_32,
        "QAG_33": q_33,
        "QAG_34": subdoc_34,
        "QAG_35": q_35,
        "QAG_36": imagens_qag_36,
        "QAG_37": subdoc_37,
        "QAG_38": q_38,
        "QAG_39": subdoc_39,
        "QAG_40": subdoc_40,
        "QAG_41": "",
        "QAG_42": "",
        "QAG_43": q_43,
        #         # "QAG_44": responsavel.data['nome'],
        #         # "QAG_45": responsavel.data['crea'],
        #         # "QAG_46": responsavel.data['ctf_ibama'],
        "QAG_47": tabela_qag47,
        "QAG_48": solventes,
        "QAG_49": imagens_qag_49,
        "QAG_54": periodicidade#PERIODICADADE SELECIONADA / AO GERAR O RELATÓRIO
        
    }

    document.render(contexto)
    document.save(local_path)

    # 8) Upload no Storage
    with open(local_path, "rb") as f:
        data = f.read()
    upload = supabase.storage.from_(BUCKET).upload(
        object_key,
        data,
        {
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        },
    )

    # 9) pegar url publica
    public_url = supabase.storage.from_(BUCKET).get_public_url(object_key)

    # 10) Insere registro na tabela `relatorios`
    try:
        supabase.table("relatorios").insert(
            {
                "nome_relatorio": payload.nome_relatorio,
                "descricao_relatorio": payload.descricao_relatorio,
                "ativo_id": str(payload.ativo_id),
                "user_id": str(payload.user_id),
                "tipo_relatorio": "qag",
                "url_relatorio": public_url,
            }
        ).execute()

    except Exception as e:
        return QAGResponse(
            sucesso=False, mensagem=f"Erro ao registrar o relatório: {str(e)}"
        )

    finally:
        # nenhum erro: devolvemos sucesso
        return QAGResponse(
            sucesso=True, mensagem="Relatório gerado e registrado com sucesso."
        )
