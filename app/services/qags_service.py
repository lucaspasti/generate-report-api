import os
import io
from uuid import UUID
from collections import OrderedDict
import datetime
from datetime import datetime
import numpy as np
import pandas as pd
import requests
from pytz import timezone
from supabase import Client
from fastapi import HTTPException
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
from app.schemas.qag import QAGRequest, QAGResponse
from app.utils.date_utils import mes_por_extenso
from app.utils.file_utils import create_signed_url
from app.utils.graficos import grafico_qualidade_agua
from app.services.indicadores.indicadores_qags import indicadores_qags
from app.services.vmps.vmp_qag import vmp_qag

parametros_inorganicos = [
    "Alumínio (mg/L)",
    "Antimônio (mg/L)",
    "Arsênio (mg/L)",
    "Bário (mg/L)",
    "Berílio (mg/L)",
    "Boro (mg/L)",
    "Cádmio (mg/L)",
    "Chumbo (mg/L)",
    "Cianeto (mg/L)",
    "Cloreto (mg/L)",
    "Cobalto (mg/L)",
    "Cobre (mg/L)",
    "Ferro (mg/L)",
    "Fluoreto (mg/L)",
    "Lítio (mg/L)",
    "Manganês (mg/L)",
    "Mercúrio (mg/L)",
    "Molibdênio (mg/L)",
    "Níquel (mg/L)",
]


parametros_organicos = [
    "Acrilamida (µg/L)",
    "Benzeno (µg/L)",
    "Benzo antraceno (µg/L)",
    "Benzo fluoranteno (µg/L)",
    "Benzo(k)fluoranteno (µg/L)",
    "Benzo pireno (µg/L)",
    "Cloreto de vinila (µg/L)",
    "Clorofórmio (µg/L)",
    "Criseno (µg/L)",
    "Dibenzo antraceno (µg/L)",
    "Diclorometano (µg/L)",
    "Estireno (µg/L)",
    "Etilbenzeno (µg/L)",
    "Fenóis (mg/L)",
]

parametros_agrotoxicos = [
    "Alaclor (µg/L)",
    "Aldicarb + ald. sulfona + ald. sulfóxido (µg/L)",
    "Aldrin + Dieldrin (µg/L)",
    "Atrazina (µg/L)",
    "Bentazona (µg/L)",
    "Carbofuran (µg/L)",
    "Clordano (cis + trans) (µg/L)",
    "Clorotalonil (µg/L)",
    "Clorpirifós 2,4-D (µg/L)",
    "Endosulfan (I + II + sulfato) (µg/L)",
    "Endrin (µg/L)",
    "Glifosato + Ampa (µg/L)",
    "Heptacloro epóxido (1024-57-3) (µg/L)",
]

parametros_microogarnismos = [
    "Coliformes fecais (NMP/100 mL)",
    "E. coli (NMP/100 mL)",
    "Enterococos (NMP/100 mL)",
]
# Caminho para o template DOCX
TEMPLATE_PATH = os.path.join(
    os.getcwd(), "app/services/relatorios/qualidade_agua_subterranea_u.docx"
)
BUCKET = "relatorios-qags"


def gerar_relatorio_qags(supabase: Client, payload: QAGRequest) -> QAGResponse:
    """
    Gera, faz upload e registra no banco um relatório QAG.
    """
    # 1) Datas e chaves
    periodicidade = payload.periodicidade
    data_str = payload.data_campanha.isoformat()
    tz_br = timezone("America/Sao_Paulo")
    now_br = datetime.now(tz_br)
    object_key = f"{payload.ativo_id}/{data_str}/{now_br:%Y-%m-%d_%H-%M-%S}.docx"
    # use o tempdir correto do sistema
    tmp_dir = tempfile.gettempdir()
    local_path = os.path.join(tmp_dir, os.path.basename(object_key))

    # 2) Carrega template
    document = DocxTemplate(TEMPLATE_PATH)

    # 3) Consultas no Supabase
    ativo = (
        supabase.table("ativos").select(
            "*").eq("id", str(payload.ativo_id)).execute()
    )
    if not ativo.data:
        raise HTTPException(404, detail="Ativo não encontrado")

    configuracoes = (
        supabase.table("configuracao_formulario_ativos")
        .select("*")
        .eq("ativo_id", str(payload.ativo_id))
        .eq("tipo_formulario", "form_qualidade_da_agua_subterranea")
        .execute()
    )
    if not configuracoes.data:
        raise HTTPException(
            404, detail="Configuração do formulário não encontrada")

    form_qags = (
        supabase.table("form_qualidade_da_agua_subterranea")
        .select("*")
        .eq("ativo_id", str(payload.ativo_id))
        .eq("campanha_de_coleta", data_str)
        .execute()
    )
    if not form_qags.data:
        raise HTTPException(404, detail="Campanha não encontrada")

    # 4) DataFrame de resultados
    resultados = form_qags.data[0]["resultados"]

    df_resultados = pd.DataFrame(resultados).fillna("Indisponível")

    parametros_escolhidos = parametros_organicos

    data_dt = datetime.strptime(data_str, "%Y-%m-%d")

    q_14 = df_resultados["Ponto"].nunique()

    q_20_1 = form_qags.data[0]["nome_laboratorio"]
    q_20_2 = form_qags.data[0]["razao_social_laboratorio"]
    q_20_3 = form_qags.data[0]["cnpj_laboratorio"]
    q_20_4 = form_qags.data[0]["endereco_laboratorio"]
    q_20_5 = form_qags.data[0]["responsavel_tecnico"]
    q_20_6 = form_qags.data[0]["email"]
    q_20_7 = form_qags.data[0]["contato"]

    url = form_qags.data[0]["registros_fotograficos_sondas"][0]

    # 2) Baixe o conteúdo
    resp = requests.get(url)
    resp.raise_for_status()  # levanta exceção se status != 200

    # reusar o mesmo buffer
    buf = io.BytesIO(resp.content)
    # ajuste a largura que quiser
    q_22 = InlineImage(document, buf, width=Cm(5))

    url = form_qags.data[0]["registros_fotograficos_amostradores"][0]

    # 2) Baixe o conteúdo
    resp = requests.get(url)
    resp.raise_for_status()  # levanta exceção se status != 200

    # reusar o mesmo buffer
    # ajuste a largura que quiser
    q_23 = InlineImage(document, buf, width=Cm(5))

    url = form_qags.data[0]["registros_fotograficos_caixas_termicas"][0]

    # 2) Baixe o conteúdo
    resp = requests.get(url)
    resp.raise_for_status()  # levanta exceção se status != 200

    # reusar o mesmo buffer
    # ajuste a largura que quiser
    q_24 = InlineImage(document, buf, width=Cm(5))

    q_25 = configuracoes.data[0]["dados_laboratoriais"][0].get(
        "metodologia_adotada")

    q_26 = indicadores_qags

    selecionados = df_resultados.copy()
    selecionados = selecionados[
        ["Ponto", "Usos Preponderantes da Água", "Profundidade", "Tipo de análise"]
        + parametros_escolhidos
    ]

    ########################################################################################################

    #############################################################################################################

    ################################################################
    q_27 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Tipo de análise"]
        + parametros_inorganicos
    ].to_dict(orient="records")

    subdoc_27 = document.new_subdoc()
    n_cols = 2 + len(
        parametros_inorganicos
    )  # Ponto + Tipo de análise + parâmetros inorgânicos
    table = subdoc_27.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Tipo de análise"
    for idx, param in enumerate(parametros_inorganicos, start=2):
        hdr[idx].text = param

    # aplicar negrito no cabeçalho
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo ordem
    grupos = OrderedDict()
    for linha in q_27:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + merge vertical na coluna Ponto
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            # preenche Tipo de análise
            row[1].text = str(linha["Tipo de análise"])
            # preenche parâmetros inorgânicos
            for j, param in enumerate(parametros_inorganicos, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1
        # merge das células da coluna Ponto
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto

    #################################################
    registros = []
    q_30 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Profundidade", "Tipo de análise"]
        + parametros_inorganicos
    ]

    for _, row in q_30.iterrows():
        classe = row["Usos Preponderantes da Água"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros.append(
                {
                    "Ponto": row["Ponto"],
                    "Usos Preponderantes da Água": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao = pd.DataFrame(registros)

    q_31 = df_comparacao["Conforme"].mean() * 100

    #####################################################

    registros_32 = []
    df_q_32 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Profundidade", "Tipo de análise"]
        + parametros_organicos
    ]

    for _, row in df_q_32.iterrows():
        classe = row["Usos Preponderantes da Água"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_32.append(
                {
                    "Ponto": row["Ponto"],
                    "Usos Preponderantes da Água": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao = pd.DataFrame(registros_32)

    q_32 = df_comparacao["Conforme"].mean() * 100

    ##############################################################################################
    q_33 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Tipo de análise"]
        + parametros_inorganicos
    ].to_dict(orient="records")

    subdoc_33 = document.new_subdoc()
    n_cols = 2 + len(
        parametros_inorganicos
    )  # Ponto + Tipo de análise + parâmetros inorgânicos
    table = subdoc_33.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Tipo de análise"
    for idx, param in enumerate(parametros_inorganicos, start=2):
        hdr[idx].text = param

    # aplicar negrito no cabeçalho
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo ordem

    grupos = OrderedDict()
    for linha in q_33:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + merge vertical na coluna Ponto
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            # preenche Tipo de análise
            row[1].text = str(linha["Tipo de análise"])
            # preenche parâmetros inorgânicos
            for j, param in enumerate(parametros_inorganicos, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1
        # merge das células da coluna Ponto
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto

    ###################################################################################################

    registros_34 = []
    df_q34 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Profundidade", "Tipo de análise"]
        + parametros_agrotoxicos
    ]

    for _, row in df_q34.iterrows():
        classe = row["Usos Preponderantes da Água"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_34.append(
                {
                    "Ponto": row["Ponto"],
                    "Usos Preponderantes da Água": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao = pd.DataFrame(registros_34)

    q_34 = df_comparacao["Conforme"].mean() * 100
    ##############################################################################
    q_35 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Tipo de análise"]
        + parametros_agrotoxicos
    ].to_dict(orient="records")

    subdoc_35 = document.new_subdoc()
    n_cols = 2 + len(
        parametros_agrotoxicos
    )  # Ponto + Tipo de análise + parâmetros inorgânicos
    table = subdoc_35.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = "Ponto"
    hdr[1].text = "Tipo de análise"
    for idx, param in enumerate(parametros_agrotoxicos, start=2):
        hdr[idx].text = param

    # aplicar negrito no cabeçalho
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # agrupa por Ponto, mantendo ordem
    grupos = OrderedDict()
    for linha in q_35:
        grupos.setdefault(linha["Ponto"], []).append(linha)

    # preenche linhas + merge vertical na coluna Ponto
    current_row = 1
    for ponto, linhas in grupos.items():
        start_row = current_row
        for linha in linhas:
            row = table.add_row().cells
            # preenche Tipo de análise
            row[1].text = str(linha["Tipo de análise"])
            # preenche parâmetros inorgânicos
            for j, param in enumerate(parametros_agrotoxicos, start=2):
                row[j].text = str(linha.get(param, ""))
            current_row += 1
        end_row = current_row - 1
        # merge das células da coluna Ponto
        cell_to_keep = table.cell(start_row, 0)
        for r in range(start_row + 1, end_row + 1):
            cell_to_keep = cell_to_keep.merge(table.cell(r, 0))
        cell_to_keep.text = ponto

    #################################################

    registros_38 = []
    df_q38 = df_resultados.copy()[
        ["Ponto", "Usos Preponderantes da Água", "Profundidade", "Tipo de análise"]
        + parametros_microogarnismos
    ]

    for _, row in df_q38.iterrows():
        classe = row["Usos Preponderantes da Água"]
        # pega o sub-dicionário de parâmetros→limites para esta classe (ou {} se não existir)
        limites = vmp_qag.get(classe, {})

        for parametro, vmp in limites.items():
            try:
                vmp = float(vmp)
                valor = float(row.get(parametro, None))
            except (ValueError, TypeError):
                valor = "Indisponível"
                vmp = "Indisponível"
            registros_38.append(
                {
                    "Ponto": row["Ponto"],
                    "Usos Preponderantes da Água": classe,
                    "Parametro": parametro,
                    "Valor": valor,
                    "VMP": vmp,
                    "Conforme": None if valor is None else (valor <= vmp),
                }
            )

    df_comparacao = pd.DataFrame(registros_38)

    q_38 = df_comparacao["Conforme"].mean() * 100
    #########################################################################

    tabela_q40 = indicadores_qags

    tabela_q40 = pd.DataFrame(tabela_q40)

    tabela_q40_aux = df_comparacao.copy()

    tabela_q40 = tabela_q40.merge(
        tabela_q40_aux[["Parametro", "Valor", "Conforme"]],
        how="left",
        left_on=["Parametro"],
        right_on=["Parametro"],
    )

    tabela_q40 = tabela_q40.dropna()

    for _, row in tabela_q40.iterrows():
        if row["Conforme"] == True:
            tabela_q40.at[_, "Conforme"] = "Alcançado"
        elif row["Conforme"] == False:
            tabela_q40.at[_, "Conforme"] = "Não Alcançado"

    tabela_q40 = tabela_q40.drop(
        columns=["Tipo", "Programa", "Valor", "Unidade", "Parametro"]
    )
    tabela_q40 = tabela_q40.rename(columns={"Conforme": "Resultado"})

    tabela_q40 = tabela_q40.to_dict(orient="records")

    #################################################################

    laudo = form_qags.data[0]["laudos"][0]

    q_43 = RichText()
    q_43.add(
        laudo, underline=True, color="#1281F0", url_id=document.build_url_id(laudo)
    )

    # @##################################################

    # 7) Contexto e renderização
    contexto = {
        "parametros_escolhidos": parametros_escolhidos,
        "QAGS_01": ativo.data[0]["nome"],
        "QAGS_02": form_qags.data[0]["campanha_de_coleta"],
        "QAGS_03": data_dt.strftime("%m"),
        "QAGS_04": data_dt.strftime("%Y"),
        "QAGS_05": "Florianópolis",
        "QAGS_06": datetime.now().day,
        "QAGS_07": mes_por_extenso(data_dt.strftime("%m")),
        "QAGS_08": datetime.now().year,
        "QAGS_09": ativo.data[0]["nome"],
        "QAGS_10": ativo.data[0]["cnpj"],
        "QAGS_11": ativo.data[0]["endereco"],
        "QAGS_12": ativo.data[0]["nome"],
        "QAGS_13": ativo.data[0]["numero_licenca"],
        "QAGS_14": q_14,
        "QAGS_15": ativo.data[0]["orgao_regulador"],
        "QAGS_16": ativo.data[0]["endereco"],
        # "QAGS_17": configuracoes.data[0]['localizacao_dos_pontos_de_monitoramento'],#pontos
        "QAGS_18": configuracoes.data[0][
            "localizacao_dos_pontos_de_monitoramento"
        ],  # pontos
        "QAGS_19": configuracoes.data[0]["parametro_periodicidade"],  # pontos
        "QAGS_20_1": q_20_1,
        "QAGS_20_2": q_20_2,
        "QAGS_20_3": q_20_3,
        "QAGS_20_4": q_20_4,
        "QAGS_20_5": q_20_5,
        "QAGS_20_6": q_20_6,
        "QAGS_20_7": q_20_7,
        #        "QAG_21": configuracoes.data[0]['parametros_periodicidade'],#
        "QAGS_22": q_22,
        "QAGS_23": q_23,
        "QAGS_24": q_24,
        "QAGS_25": q_25,
        "QAGS_26": q_26,
        "tabela_qags_27": subdoc_27,
        "QAGS_28": parametros_inorganicos,
        #        "QAGS_29": imagens_qag29,
        "QAGS_31": q_31,
        "QAGS_32": q_32,
        "QAGS_33": subdoc_33,
        "QAGS_34": q_34,
        "QAGS_35": subdoc_35,
        #        "QAG_36": ,
        #         "QAG_37": subdoc_37,
        "QAG_38": q_38,
        #         "QAG_39": subdoc_39,
        "QAG_40": tabela_q40,
        #         "QAG_41": '',
        #         "QAG_42": '',
        "QAG_43": q_43,
        "QAG_54": periodicidade  # PERIODICADADE SELECIONADA / AO GERAR O RELATÓRIO

    }

    document.render(contexto)
    document.save(local_path)

    # 8) Upload no Storage
    with open(local_path, "rb") as f:
        data = f.read()
    upload = supabase.storage.from_(BUCKET).upload(
        object_key,
        data,
        {
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        },
    )

    # 9) pegar url publica
    public_url = supabase.storage.from_(BUCKET).get_public_url(object_key)

    # 10) Insere registro na tabela `relatorios`
    try:
        supabase.table("relatorios").insert(
            {
                "nome_relatorio": payload.nome_relatorio,
                "descricao_relatorio": payload.descricao_relatorio,
                "ativo_id": str(payload.ativo_id),
                "user_id": str(payload.user_id),
                "tipo_relatorio": "qag",
                "url_relatorio": public_url,
            }
        ).execute()

    except Exception as e:
        return QAGResponse(
            sucesso=False, mensagem=f"Erro ao registrar o relatório: {str(e)}"
        )

    finally:
        # nenhum erro: devolvemos sucesso
        return QAGResponse(
            sucesso=True, mensagem="Relatório gerado e registrado com sucesso."
        )
